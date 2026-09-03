"""DOCX file ingestor."""

from __future__ import annotations

from pathlib import Path

from app.domain.documents import DocumentMetadata, SourceDocument
from app.infrastructure.ingestion.base import (
    BaseIngestor,
    IngestionError,
    SourceReference,
    require_path_source,
)
from app.infrastructure.ingestion.utils import file_timestamp


class DocxIngestor(BaseIngestor):
    """Read DOCX files into normalized source documents."""

    source_type = "docx"
    supported_suffixes = (".docx", ".odt")

    def ingest(self, source: SourceReference) -> SourceDocument:
        source_path = require_path_source(source, ingestor_name="DOCX ingestor")
        text = self._extract_text(source_path)
        resolved_path = source_path.resolve()
        is_odt = source_path.suffix.lower() == ".odt"
        mime = (
            "application/vnd.oasis.opendocument.text"
            if is_odt
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return SourceDocument(
            source=str(resolved_path),
            source_path=resolved_path,
            source_type=self.source_type,
            filename=source_path.name,
            text=text,
            metadata=DocumentMetadata(
                title=source_path.stem,
                modified_at=file_timestamp(source_path),
                mime_type=mime,
            ),
        )

    def _extract_text(self, path: Path) -> str:
        if path.suffix.lower() == ".odt":
            return self._extract_odt(path)
        return self._extract_docx(path)

    def _extract_docx(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore[import-untyped]

            doc = Document(str(path))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except ImportError:
            raise IngestionError(
                "python-docx is required for DOCX ingestion. Install with: pip install python-docx"
            ) from None
        except Exception as exc:
            raise IngestionError(f"Unable to read DOCX file '{path}'.") from exc

    def _extract_odt(self, path: Path) -> str:
        try:
            from odf import teletype
            from odf import text as odf_text  # type: ignore[import-untyped]
            from odf.opendocument import load  # type: ignore[import-untyped]

            doc = load(str(path))
            parts = [
                teletype.extractText(p)
                for p in doc.getElementsByType(odf_text.P)
                if teletype.extractText(p).strip()
            ]
            return "\n".join(parts)
        except ImportError:
            raise IngestionError(
                "odfpy is required for ODT ingestion. Install with: pip install odfpy"
            ) from None
        except IngestionError:
            raise
        except Exception as exc:
            raise IngestionError(f"Unable to read ODT file '{path}'.") from exc
